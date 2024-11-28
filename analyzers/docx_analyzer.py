from docx import Document
from typing import Dict, Any, List, Tuple
from collections import Counter
from .base_analyzer import BaseAnalyzer


class DocxAnalyzer(BaseAnalyzer):
    """Анализатор DOCX файлов."""

    def analyze(self) -> Dict[str, Any]:
        """Анализ DOCX файла."""
        if not self.validate_file():
            return {"error": "Ошибка при открытии файла"}

        try:
            doc = Document(self.file_path)
        except Exception as e:
            self.logger.error(f"Ошибка при открытии DOCX файла: {str(e)}")
            return {"error": f"Ошибка при открытии DOCX файла: {str(e)}"}

        try:
            # Базовая статистика
            stats = self._collect_basic_stats(doc)
            
            # Анализ форматирования
            formatting = self._analyze_formatting(doc)
            
            # Анализ таблиц
            table_stats = self._analyze_tables(doc)
            
            # Анализ стилей
            style_stats = self._analyze_styles(doc)

            results = {
                "Общая статистика": [
                    f"Количество абзацев: {stats['paragraphs']}",
                    f"Количество строк (runs): {stats['runs']}",
                    f"Количество слов: {stats['words']}",
                    f"Количество символов: {stats['chars']}",
                    f"Количество заголовков: {stats['headings']}",
                    f"Количество изображений: {stats['images']}",
                    f"Количество таблиц: {stats['tables']}"
                ],
                "Форматирование текста": [
                    f"Жирный текст: {formatting['bold']}",
                    f"Курсивный текст: {formatting['italic']}",
                    f"Подчеркнутый текст: {formatting['underline']}"
                ],
                "Статистика таблиц": [
                    f"Всего ячеек: {table_stats['total_cells']}",
                    f"Среднее количество строк: {table_stats['avg_rows']:.1f}",
                    f"Среднее количество столбцов: {table_stats['avg_cols']:.1f}",
                    f"Пустые ячейки: {table_stats['empty_cells']}"
                ],
                "Статистика стилей": [
                    f"{style}: {count}" for style, count in style_stats.items()
                ]
            }

            return results

        except Exception as e:
            self.logger.error(f"Ошибка при анализе DOCX файла: {str(e)}")
            return {"error": f"Ошибка при анализе DOCX файла: {str(e)}"}

    def _collect_basic_stats(self, doc: Document) -> Dict[str, int]:
        """Сбор базовой статистики документа."""
        stats = {
            'paragraphs': len(doc.paragraphs),
            'runs': 0,
            'words': 0,
            'chars': 0,
            'headings': 0,
            'images': 0,
            'tables': len(doc.tables)
        }

        for paragraph in doc.paragraphs:
            stats['runs'] += len(paragraph.runs)
            stats['words'] += len(paragraph.text.split())
            stats['chars'] += len(paragraph.text)

            if paragraph.style.name.startswith('Heading'):
                stats['headings'] += 1

        # Подсчет изображений
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                stats['images'] += 1

        return stats

    def _analyze_formatting(self, doc: Document) -> Dict[str, int]:
        """Анализ форматирования текста."""
        formatting = {
            'bold': 0,
            'italic': 0,
            'underline': 0
        }

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    formatting['bold'] += 1
                if run.italic:
                    formatting['italic'] += 1
                if run.underline:
                    formatting['underline'] += 1

        return formatting

    def _analyze_tables(self, doc: Document) -> Dict[str, float]:
        """Анализ таблиц в документе."""
        if not doc.tables:
            return {
                'total_cells': 0,
                'avg_rows': 0,
                'avg_cols': 0,
                'empty_cells': 0
            }

        total_rows = 0
        total_cols = 0
        total_cells = 0
        empty_cells = 0

        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns)
            total_rows += rows
            total_cols += cols
            total_cells += rows * cols

            for row in table.rows:
                for cell in row.cells:
                    if not cell.text.strip():
                        empty_cells += 1

        num_tables = len(doc.tables)
        return {
            'total_cells': total_cells,
            'avg_rows': total_rows / num_tables if num_tables > 0 else 0,
            'avg_cols': total_cols / num_tables if num_tables > 0 else 0,
            'empty_cells': empty_cells
        }

    def _analyze_styles(self, doc: Document) -> Dict[str, int]:
        """Анализ использования стилей в документе."""
        style_counter = Counter()
        
        for paragraph in doc.paragraphs:
            style_counter[paragraph.style.name] += 1

        return dict(style_counter.most_common())


def analyze_docx(file_path: str) -> str:
    """Функция-обертка для обратной совместимости."""
    analyzer = DocxAnalyzer(file_path)
    results = analyzer.analyze()
    return analyzer.format_results(results)
